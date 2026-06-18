import re
import logging
from pathlib import Path
from pypdf import PdfReader
import docx

logger = logging.getLogger(__name__)

def parse_pdf(file_path):
    """
    Reads text from a PDF file using pypdf and parses questions.
    """
    try:
        reader = PdfReader(file_path)
        all_text = []
        for idx, page in enumerate(reader.pages):
            text = page.extract_text()
            if not text:
                continue
            all_text.append(f"--- SAYFA {idx + 1} ---")
            all_text.append(text)
        
        full_text = "\n".join(all_text)
        return parse_questions_from_text(full_text)
    except Exception as e:
        logger.error(f"PDF okuma hatası: {e}")
        raise e

def parse_docx(file_path):
    """
    Reads text from a Word (.docx) file using python-docx and parses questions.
    """
    try:
        doc = docx.Document(file_path)
        text_parts = []
        
        for p in doc.paragraphs:
            text_parts.append(p.text)
            
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = "\n".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                    if not cell_text:
                        continue
                    row_text.append(cell_text)
                if not row_text:
                    continue
                text_parts.append(" | ".join(row_text))
                
        full_text = "\n".join(text_parts)
        return parse_questions_from_text(full_text)
    except Exception as e:
        logger.error(f"Word okuma hatası: {e}")
        raise e

def _extract_tab_options(line):
    """
    Tab ile ayrılmış şık formatını ayrıştır.
    Örnek: '\tA). A2\tB). 2A\tC). Hücre\tD). Fonksiyon\tE). Değer'
    veya: '\tA) Seçenek\tB) Seçenek2...'
    """
    # A). / A) / A. formatlarını tab ile ayrılmış olarak yakala
    parts = re.split(r'\t', line)
    options = {}
    for part in parts:
        part = part.strip()
        m = re.match(r'^([A-Ea-e])[\s\)\.\-]+\s*(.*)', part)
        if m:
            letter = m.group(1).lower()
            content = m.group(2).strip()
            if content:
                options[letter] = content
    return options

def parse_questions_from_text(text):
    """
    Common parser that segments raw text into question objects using regex patterns.
    Desteklenen formatlar:
      - '1. Soru metni' veya '1) Soru metni'
      - 'S1. Soru metni' veya 'S1) Soru metni'  (S öneki)
      - Şıklar: 'A) ...' veya 'A). ...' veya tab ayrılmış '\tA). ...\tB). ...'
    """
    lines = [line for line in text.split("\n")]
    questions = []
    current_q = None
    
    # Soru başlangıç patternleri: 1. / 1) / S1. / S1)
    q_start_pat = re.compile(r'^(?:S|s)?(\d+)[\.\)]\s+(.+)$')
    # Şık pattern: A) / A). / A- / A. 
    opt_pat = re.compile(r'^([A-Ea-e])[\s\)\.\-⪢]+\s*(.*)$')
    # Satır içi çok şık (tab ayrılmış)
    inline_opt_pat = re.compile(r'\t([A-Ea-e])[\)\.\-\s]+\s*')
    
    for line in lines:
        stripped = line.strip()
        
        if not stripped:
            continue
        if stripped.startswith("--- SAYFA") or stripped == "Diğer sayfaya geçiniz." or \
           stripped.startswith("___") or stripped.startswith("---"):
            continue
        
        # Tab ile ayrılmış satır ise şık mı diye kontrol et
        if '\t' in line and current_q is not None:
            tab_opts = _extract_tab_options(line)
            if tab_opts:
                for letter, content in tab_opts.items():
                    current_q['options'][letter] = content
                    current_q['current_opt'] = letter
                continue
        
        # Soru başlangıcı
        match_q = q_start_pat.match(stripped)
        if match_q:
            if current_q:
                questions.append(current_q)
            q_num = match_q.group(1)
            q_text = match_q.group(2).strip()
            current_q = {
                'number': q_num,
                'text': q_text,
                'options': {'a': '', 'b': '', 'c': '', 'd': '', 'e': ''},
                'correct_answer': 'A',
                'current_opt': None
            }
            continue
            
        if current_q is None:
            continue

        # Satır içi birden fazla şık (tab yok ama regex ile)
        inline_matches = list(re.finditer(r'\b([A-Ea-e])[\s\)\.\-⪢]+\s*', stripped))
        if len(inline_matches) >= 2:
            all_letters = {m.group(1).lower() for m in inline_matches}
            if all_letters & {'a', 'b', 'c', 'd'}:
                for i in range(len(inline_matches)):
                    opt_letter = inline_matches[i].group(1).lower()
                    if opt_letter not in ('a', 'b', 'c', 'd', 'e'):
                        continue
                    start_idx = inline_matches[i].end()
                    end_idx = inline_matches[i + 1].start() if i + 1 < len(inline_matches) else len(stripped)
                    content = stripped[start_idx:end_idx].strip()
                    if content:
                        current_q['options'][opt_letter] = content
                        current_q['current_opt'] = opt_letter
                continue
            
        # Tek şık
        match_opt = opt_pat.match(stripped)
        if match_opt:
            opt_letter = match_opt.group(1).lower()
            content = match_opt.group(2).strip()
            current_q['options'][opt_letter] = content
            current_q['current_opt'] = opt_letter
            continue
            
        # Devam satırı
        if current_q['current_opt']:
            existing = current_q['options'].get(current_q['current_opt'], '')
            if existing:
                current_q['options'][current_q['current_opt']] += ' ' + stripped
            else:
                current_q['options'][current_q['current_opt']] = stripped
        else:
            if current_q['text']:
                current_q['text'] += '\n' + stripped
            else:
                current_q['text'] = stripped
                
    if current_q:
        questions.append(current_q)
        
    valid_questions = []
    for q in questions:
        for k in q['options']:
            q['options'][k] = q['options'][k].strip()
        q['text'] = q['text'].strip()
        
        filled_opts = sum(1 for v in q['options'].values() if v.strip())
        if q['text'] and filled_opts >= 2:
            if 'current_opt' in q:
                del q['current_opt']
            valid_questions.append(q)
            
    return valid_questions
