import logging
from pathlib import Path
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import config

logger = logging.getLogger(__name__)

class WordGenerator:
    """Word belgesi oluşturma sınıfı"""

    @staticmethod
    def create_exam(exam_data, questions, output_path):
        try:
            import shutil
            taslak_path = config.BASE_DIR / 'taslak.docx'
            if taslak_path.exists():
                shutil.copy(str(taslak_path), output_path)
                doc = Document(output_path)
                for _ in range(3):
                    doc.add_paragraph()
            else:
                doc = Document()
                section = doc.sections[0]
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
                
                template_path = config.TEMPLATES_DIR / 'omr_template.png'
                if template_path.exists():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(template_path), width=Inches(5.5))
                    for _ in range(3):
                        doc.add_paragraph()
                else:
                    section.top_margin = Inches(1.0)
                    section.bottom_margin = Inches(1.0)
                    section.left_margin = Inches(1.0)
                    section.right_margin = Inches(1.0)
                    
                    WordGenerator._add_header(doc, exam_data)
                    WordGenerator._add_instructions(doc, exam_data)
                    WordGenerator._add_answer_sheet(doc, exam_data, questions)

            from database import db_manager
            layout_cols = int(db_manager.get_setting('layout_columns', 2))
            
            if layout_cols > 1:
                for section in doc.sections:
                    sectPr = section._sectPr
                    cols = sectPr.xpath('./w:cols')
                    if not cols:
                        cols = OxmlElement('w:cols')
                        sectPr.append(cols)
                    else:
                        cols = cols[0]
                    cols.set(qn('w:num'), str(layout_cols))
                    cols.set(qn('w:space'), '360')
                    
            WordGenerator._add_questions(doc, questions, exam_data)
            doc.save(output_path)
            logger.info(f"Sınav belgesi oluşturuldu: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Word belgesi oluşturma hatası: {e}")
            raise e

    @staticmethod
    def _add_header(doc, exam_data):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(config.APP_AUTHOR)
        run.font.size = Pt(14)
        run.font.bold = True
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(exam_data.get('course_name', ''))
        run.font.size = Pt(13)
        run.font.bold = True
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        exam_type = exam_data.get('exam_type', '')
        exam_group = exam_data.get('exam_group', '')
        exam_date = exam_data.get('exam_date', '')
        
        run = p.add_run(f"{exam_type} Sınavı - Grup {exam_group}")
        run.font.size = Pt(12)
        run.font.bold = True
        
        if exam_date:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Tarih: {exam_date}")
            run.font.size = Pt(11)
            
        table = doc.add_table(rows=2, cols=3)
        table.style = 'Table Grid'
        
        headers = ['Ad Soyad:', 'Öğrenci No:', 'İmza:']
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            
        for i in range(3):
            table.cell(1, i).text = ''
            
        doc.add_paragraph()

    @staticmethod
    def _add_instructions(doc, exam_data):
        p = doc.add_paragraph()
        run = p.add_run('AÇIKLAMALAR:')
        run.bold = True
        run.font.size = Pt(11)
        
        instructions = [
            f"Sınav süresi: {exam_data.get('duration', 60)} dakika",
            f"Toplam soru sayısı: {exam_data.get('question_count', 0)}",
            "Cevaplarınızı sınav kağıdına işaretleyiniz.",
            "Her sorunun sadece bir doğru cevabı vardır.",
            "Başarılar dileriz."
        ]
        
        for instruction in instructions:
            p = doc.add_paragraph(instruction, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            
        doc.add_paragraph()

    @staticmethod
    def _add_answer_sheet(doc, exam_data, questions):
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)
        
        right_par = right_cell.paragraphs[0]
        right_par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = right_par.add_run(
            f"SINAV KAĞIDI   {exam_data.get('course_name', '')} - {exam_data.get('exam_type', '')}"
        )
        run.font.size = Pt(11)
        run.font.bold = True
        
        meta_lines = [
            f"Dersin Sorumlusu: {exam_data.get('instructor', '')}",
            f"Sınav Tarihi: {exam_data.get('exam_date', '')}",
            f"Sınav Süresi: {exam_data.get('duration', '')} dk"
        ]
        
        for line in meta_lines:
            p = right_cell.add_paragraph(line)
            p.paragraph_format.space_before = Pt(2)
            p.runs[0].font.size = Pt(9)
            
        logo_path = config.TEMPLATES_DIR / 'logo.png'
        if logo_path.exists():
            try:
                left_cell.paragraphs[0].clear()
                lp = left_cell.paragraphs[0]
                r = lp.add_run()
                r.add_picture(str(logo_path), width=Inches(1.4))
            except Exception:
                pass
                
        total_q = len(questions) if questions else 60
        total_q = max(total_q, 60)
        main_rows = 34
        cols = 6
        
        omr_table = left_cell.add_table(rows=main_rows, cols=cols)
        omr_table.style = 'Table Grid'
        bubble = '○'
        
        for i in range(main_rows):
            qn_val = i + 1
            cell = omr_table.cell(i, 0)
            cell.text = str(qn_val)
            para = cell.paragraphs[0]
            if para.runs:
                para.runs[0].font.size = Pt(9)
            else:
                para.add_run(str(qn_val)).font.size = Pt(9)
                
            for j in range(1, cols):
                c = omr_table.cell(i, j)
                c.text = bubble
                para = c.paragraphs[0]
                if para.runs:
                    para.runs[0].font.size = Pt(12)
                else:
                    para.add_run(bubble).font.size = Pt(12)
                    
        left_cell.add_paragraph()
        small_table = left_cell.add_table(rows=26, cols=2)
        small_table.style = 'Table Grid'
        
        for idx in range(26):
            qnum = 35 + idx
            r = small_table.rows[idx]
            r.cells[0].text = str(qnum)
            para0 = r.cells[0].paragraphs[0]
            if para0.runs:
                para0.runs[0].font.size = Pt(9)
            else:
                para0.add_run(str(qnum)).font.size = Pt(9)
                
            bubble_str = f"  {bubble} {bubble} {bubble} {bubble} {bubble}"
            r.cells[1].text = bubble_str
            para1 = r.cells[1].paragraphs[0]
            if para1.runs:
                para1.runs[0].font.size = Pt(11)
            else:
                para1.add_run(bubble_str).font.size = Pt(11)
                
        right_cell.add_paragraph()
        qlist_title = right_cell.add_paragraph('Soru Listesi')
        qlist_title.runs[0].font.bold = True
        qlist_title.runs[0].font.size = Pt(10)
        
        for idx, q in enumerate(questions[:30]):
            qnum = idx + 1
            qtext = q.get('question_text', '')
            short = qtext[:120] + '...' if len(qtext) > 120 else qtext
            p = right_cell.add_paragraph()
            run = p.add_run(f"{qnum}) {short}")
            run.font.size = Pt(9)
            
        doc.add_paragraph()

    @staticmethod
    def _add_questions(doc, questions, exam_data):
        for idx, question in enumerate(questions, 1):
            question_paragraphs = []
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.05
            
            run = p.add_run(f"{idx}. ")
            run.font.bold = True
            run.font.size = Pt(11)
            
            run = p.add_run(question.get('question_text', ''))
            run.font.bold = True
            run.font.size = Pt(11)
            question_paragraphs.append(p)
            
            question_image = question.get('question_image_path')
            if question_image and Path(question_image).exists():
                try:
                    p_img = doc.add_paragraph()
                    p_img.paragraph_format.space_before = Pt(2)
                    p_img.paragraph_format.space_after = Pt(2)
                    p_img.paragraph_format.left_indent = Inches(0.4)
                    run_img = p_img.add_run()
                    width_val = config.EXAM_SETTINGS['question_image_width'] / 100
                    run_img.add_picture(question_image, width=Inches(width_val))
                    question_paragraphs.append(p_img)
                except Exception as e:
                    logger.warning(f"Soru resmi eklenemedi: {e}")
                    
            options_to_write = []
            for option_char in ('a', 'b', 'c', 'd', 'e'):
                opt_text = question.get(f'option_{option_char}')
                opt_image = question.get(f'option_{option_char}_image_path')
                if opt_text or (opt_image and Path(opt_image).exists()):
                    options_to_write.append({
                        'char': option_char.upper(),
                        'text': opt_text if opt_text else '',
                        'image': opt_image
                    })
                    
            lines = []
            current_line = []
            current_len = 0
            
            from database import db_manager
            LINE_LIMIT = int(db_manager.get_setting('option_wrap_limit', 40))
            
            for opt in options_to_write:
                if opt['image'] and Path(opt['image']).exists():
                    if current_line:
                        lines.append(current_line)
                        current_line = []
                        current_len = 0
                    lines.append([opt])
                else:
                    opt_str = f"{opt['char']}) {opt['text']}"
                    opt_len = len(opt_str) + 8
                    if not current_line:
                        current_line.append(opt)
                        current_len = opt_len
                    elif current_len + opt_len <= LINE_LIMIT:
                        current_line.append(opt)
                        current_len += opt_len
                    else:
                        lines.append(current_line)
                        current_line = [opt]
                        current_len = opt_len
            if current_line:
                lines.append(current_line)
                
            for line in lines:
                p_opt = doc.add_paragraph()
                p_opt.paragraph_format.left_indent = Inches(0.4)
                p_opt.paragraph_format.space_before = Pt(0)
                p_opt.paragraph_format.space_after = Pt(2)
                p_opt.paragraph_format.line_spacing = 1.05
                
                for i, opt in enumerate(line):
                    if i > 0:
                        p_opt.add_run('\t')
                    run_opt_char = p_opt.add_run(f"{opt['char']}) ")
                    run_opt_char.font.bold = True
                    run_opt_char.font.size = Pt(11)
                    
                    if opt['text']:
                        run_opt_text = p_opt.add_run(opt['text'])
                        run_opt_text.font.size = Pt(11)
                        
                    if opt['image'] and Path(opt['image']).exists():
                        try:
                            run_opt_img = p_opt.add_run()
                            width_val = config.EXAM_SETTINGS['answer_image_width'] / 100
                            run_opt_img.add_picture(opt['image'], width=Inches(width_val))
                        except Exception as e:
                            logger.warning(f"Şık resmi eklenemedi: {e}")
                question_paragraphs.append(p_opt)
                
            for p_item in question_paragraphs[:-1]:
                p_item.paragraph_format.keep_with_next = True

    @staticmethod
    def create_answer_key(exam_data, questions, output_path):
        try:
            doc = Document()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('CEVAP ANAHTARI')
            run.font.size = Pt(14)
            run.font.bold = True
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(
                f"{exam_data.get('course_name', '')} - {exam_data.get('exam_type', '')} - Grup {exam_data.get('exam_group', '')}"
            )
            run.font.size = Pt(12)
            
            doc.add_paragraph()
            rows_needed = (len(questions) + 4) // 5
            table = doc.add_table(rows=rows_needed + 1, cols=10)
            table.style = 'Table Grid'
            
            headers = []
            for i in range(5):
                headers.extend(['Soru', 'Cevap'])
                
            for col_idx, header in enumerate(headers):
                cell = table.cell(0, col_idx)
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            for idx, question in enumerate(questions):
                row = (idx // 5) + 1
                col_offset = (idx % 5) * 2
                
                table.cell(row, col_offset).text = str(idx + 1)
                table.cell(row, col_offset).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                answer = question.get('correct_answer', '-')
                cell = table.cell(row, col_offset + 1)
                cell.text = answer
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].runs[0].font.bold = True
                
            doc.save(output_path)
            logger.info(f"Cevap anahtarı oluşturuldu: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Cevap anahtarı oluşturma hatası: {e}")
            raise e
